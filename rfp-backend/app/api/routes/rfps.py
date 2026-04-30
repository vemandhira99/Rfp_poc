from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.api.routes.auth import get_current_user
from app.services import rfp_service
from app.schemas.rfp import RFPOut, DecisionRequest, AssignRequest, CommentRequest, DashboardSummary, ChatRequest, DraftRequest, RequirementOut, NotificationOut
from datetime import datetime
from typing import List
from app.models.rfp import RFPDocument, RFPDraft, AuditLog

router = APIRouter(prefix="/rfps", tags=["RFPs"])

@router.get("/dashboard-summary")
def dashboard_summary(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    return rfp_service.get_dashboard_summary(db)

@router.get("/", response_model=List[RFPOut])
def list_rfps(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    return rfp_service.get_all_rfps(db)

@router.get("/assigned-to-me", response_model=List[RFPOut])
def assigned_to_me(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    return rfp_service.get_assigned_rfps(db, current_user.id)

@router.get("/notifications", response_model=List[NotificationOut])
def get_notifications(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    from app.services import notification_service
    return notification_service.get_user_notifications(db, current_user.id)

@router.post("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    from app.services import notification_service
    notification_service.mark_as_read(db, notification_id)
    return {"message": "Notification marked as read"}

@router.post("/notifications/clear-all")
def clear_all_notifications(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    from app.services import notification_service
    notification_service.clear_all(db, current_user.id)
    return {"message": "All notifications marked as read"}

@router.get("/{rfp_id}", response_model=RFPOut)
def get_rfp(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    rfp = rfp_service.get_rfp_by_id(db, rfp_id)
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    return rfp

@router.post("/{rfp_id}/decision")
def make_decision(rfp_id: int, request: DecisionRequest,
                  db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    if current_user.role not in ["CEO", "Admin", "PM", "Leadership"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    result = rfp_service.save_decision(db, rfp_id, current_user.id, request.decision, request.reason)
    return {"message": f"Decision '{request.decision}' saved", "rfp_id": rfp_id}

def background_draft_and_compliance(rfp_id: int):
    from app.core.database import SessionLocal
    from app.services.ai_service import generate_architect_draft, extract_compliance_matrix
    db = SessionLocal()
    try:
        rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
        if rfp:
            rfp.current_status = "in_drafting"
            db.commit()

        # 1. Generate Draft
        result = generate_architect_draft(db, rfp_id)
        
        # If cancelled, stop the whole process
        if isinstance(result, dict) and result.get("status") == "cancelled":
            print(f"Background task for RFP {rfp_id} stopped due to cancellation.")
            return

        # 2. Extract Compliance Matrix
        extract_compliance_matrix(db, rfp_id)

        # Final refresh to check if user cancelled during compliance extraction
        db.refresh(rfp)
        if rfp.current_status == "in_drafting":
            rfp.current_status = "assigned_to_sa"
            db.commit()
    except Exception as e:
        print(f"Background draft/compliance failed for RFP {rfp_id}: {str(e)}")
    finally:
        db.close()

@router.post("/{rfp_id}/assign-architect")
def assign_architect(rfp_id: int, request: AssignRequest, background_tasks: BackgroundTasks,
                      db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    if current_user.role not in ["CEO", "PM", "Admin"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    result = rfp_service.assign_architect(db, rfp_id, request.architect_id, current_user.id, request.notes)
    
    # Trigger background drafting
    background_tasks.add_task(background_draft_and_compliance, rfp_id)
    
    return {"message": "Architect assigned and draft generation started", "rfp_id": rfp_id}

@router.post("/{rfp_id}/comment")
def add_comment(rfp_id: int, request: CommentRequest,
                db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    result = rfp_service.add_comment(db, rfp_id, current_user.id,
                                     request.comment_text, request.entity_type, request.entity_id)
    return {"message": "Comment added", "comment_id": result.id}

from app.services import draft_service
from app.schemas.rfp import DraftCreate, DraftOut

@router.post("/{rfp_id}/draft", response_model=DraftOut)
def save_draft(rfp_id: int, request: DraftCreate,
              db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    if current_user.role not in ["Solution_Architect", "CEO", "Admin"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    draft = draft_service.save_draft(db, rfp_id, current_user.id, request.draft_content, request.is_final)
    return draft

@router.get("/{rfp_id}/draft", response_model=DraftOut)
def get_draft(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    draft = draft_service.get_latest_draft(db, rfp_id)
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    return draft

@router.post("/{rfp_id}/chat-stream")
def chat_with_rfp_stream(rfp_id: int, request: ChatRequest, db: Session = Depends(get_db)):
    try:
        from app.services.ai_service import stream_chat_with_document
        return StreamingResponse(
            stream_chat_with_document(
                db=db, 
                rfp_id=rfp_id, 
                message=request.message, 
                knowledge_mode=request.knowledge_mode,
                history=request.history
            ),
            media_type="text/event-stream"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{rfp_id}/export")
def export_rfp(rfp_id: int, db: Session = Depends(get_db)):
    rfp = rfp_service.get_rfp_by_id(db, rfp_id)
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    # Collect all draft sections for the latest version
    from sqlalchemy import desc
    latest_v_record = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id).order_by(desc(RFPDraft.version)).first()
    if not latest_v_record:
        content = "# No Draft Available\n\nPlease generate a draft first."
    else:
        version = latest_v_record.version
        drafts = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id, RFPDraft.version == version).order_by(RFPDraft.section_order.asc()).all()
        content = ""
        for d in drafts:
            if d.section_name:
                content += f"\n\n# {d.section_name}\n\n{d.draft_content or ''}"
    
    import os, time, re
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import FileResponse
    from app.core.config import settings
    from app.models.rfp import RFPRequirement

    doc = Document()
    
    # 1. Professional Title Page
    title_p = doc.add_heading(f"\nRFP Response Draft\n\n{rfp.title}", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Client: {rfp.client_name or 'N/A'}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(f"Generated Date: {time.strftime('%Y-%m-%d')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 2. Dynamic Table of Contents
    doc.add_heading("Table of Contents", level=1)
    if latest_v_record:
        all_sections = db.query(RFPDraft).filter(
            RFPDraft.rfp_id == rfp_id, 
            RFPDraft.version == latest_v_record.version,
            RFPDraft.section_name != None
        ).order_by(RFPDraft.section_order.asc()).all()
        for s in all_sections:
            doc.add_paragraph(f"{s.section_order}. {s.section_name}")
    else:
        doc.add_paragraph("1. Executive Summary")
    doc.add_page_break()

    # 3. Parse Markdown Content into Word Sections
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('* '): doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('- '): doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line): doc.add_paragraph(line, style='List Number')
        else: doc.add_paragraph(line)

    doc.add_page_break()

    # 4. Compliance Matrix Section
    doc.add_heading("Annexure: Compliance Matrix", level=1)
    requirements = db.query(RFPRequirement).filter(RFPRequirement.rfp_id == rfp_id).all()
    if requirements:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'; hdr_cells[1].text = 'Requirement'; hdr_cells[2].text = 'Status'; hdr_cells[3].text = 'Response Strategy'
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for r in p.runs: r.bold = True
        for req in requirements:
            row_cells = table.add_row().cells
            row_cells[0].text = str(req.category or "General")
            row_cells[1].text = str(req.requirement_text)
            row_cells[2].text = str(req.status or "Pending")
            row_cells[3].text = str(req.response_strategy or "N/A")
    
    # 5. Save and Export
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    safe_title = re.sub(r'[^\w\s-]', '', rfp.title).replace(' ', '_')
    file_name = f"Response_Draft_{safe_title}_{int(time.time())}.docx"
    file_path = os.path.join(settings.UPLOAD_DIR, file_name)
    doc.save(file_path)
    return FileResponse(file_path, filename=file_name)

import json
from app.services.ai_service import chat_with_document

@router.get("/{rfp_id}/summary")
def get_ai_summary(rfp_id: int, db: Session = Depends(get_db)):
    log = db.query(AuditLog).filter(AuditLog.rfp_id == rfp_id, AuditLog.action == "ai_summary_generated").order_by(AuditLog.created_at.desc()).first()
    if not log or not log.new_value: return {"error": "Summary not found"}
    try: return json.loads(log.new_value)
    except: return {"error": "Invalid summary format"}

@router.post("/{rfp_id}/chat")
def chat_with_rfp(rfp_id: int, request: ChatRequest, db: Session = Depends(get_db)):
    try:
        reply = chat_with_document(db=db, rfp_id=rfp_id, message=request.message, knowledge_mode=request.knowledge_mode, history=request.history)
        return {"reply": reply}
    except Exception as e:
        return {"reply": f"AI Advisor is currently unavailable. ({str(e)})"}

@router.get("/{rfp_id}/compliance", response_model=List[RequirementOut])
def get_compliance_matrix(rfp_id: int, db: Session = Depends(get_db)):
    from app.models.rfp import RFPRequirement
    return db.query(RFPRequirement).filter(RFPRequirement.rfp_id == rfp_id).all()

@router.post("/{rfp_id}/regenerate")
def regenerate_proposal(rfp_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    # Reset status to in_drafting to start the loop
    from app.services import rfp_service
    rfp_service.update_rfp_status(db, rfp_id, "in_drafting", current_user.id)
    background_tasks.add_task(background_draft_and_compliance, rfp_id)
    return {"message": "Regeneration started", "status": "processing"}

@router.post("/{rfp_id}/cancel-generation")
def cancel_generation(rfp_id: int, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    from app.services import rfp_service
    # Setting status to 'on_hold' will trigger the loop in ai_service to exit
    rfp_service.update_rfp_status(db, rfp_id, "on_hold", current_user.id)
    return {"message": "Generation cancellation requested", "status": "cancelling"}

@router.get("/{rfp_id}/generation-progress")
def get_generation_progress(rfp_id: int, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    from sqlalchemy import desc
    latest_v_record = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id).order_by(desc(RFPDraft.version)).first()
    
    # If no record at all, it hasn't started
    if not latest_v_record: return {"current": 0, "total": 21, "status": "starting"}
    
    latest_version = latest_v_record.version
    count = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id, RFPDraft.version == latest_version, RFPDraft.section_name != None).count()
    
    # Check RFP document for current status (to handle manual stops/cancellations)
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    current_status = rfp.current_status if rfp else "unknown"
    
    # Determine progress status
    if count >= 21:
        status = "completed"
    elif current_status == "on_hold":
        status = "cancelled"
    elif current_status == "error" or current_status == "rate_limit_error":
        status = "failed"
    else:
        status = "processing"
        
    return {"current": count, "total": 21, "version": latest_version, "status": status}
